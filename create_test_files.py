#!/usr/bin/env python3
"""
Test file generator for Project Sentinel
Creates various safe test files to validate the analysis system
"""

import os
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import tempfile

def create_test_directory():
    """Create a test_files directory"""
    test_dir = "test_files"
    if not os.path.exists(test_dir):
        os.makedirs(test_dir)
    return test_dir

def create_simple_pdf(test_dir):
    """Create a simple PDF with some JavaScript for form validation"""
    filename = os.path.join(test_dir, "simple_form.pdf")
    
    # Create a simple PDF with some JavaScript
    c = canvas.Canvas(filename, pagesize=letter)
    
    # Add some content
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "Test PDF Document")
    c.drawString(100, 730, "This is a simple PDF for testing purposes.")
    c.drawString(100, 710, "It contains basic form validation JavaScript.")
    
    # Add some JavaScript for form validation (legitimate use)
    js_code = """
    function validateForm() {
        var x = document.forms["myForm"]["fname"].value;
        if (x == "") {
            alert("Name must be filled out");
            return false;
        }
        return true;
    }
    """
    
    # Note: ReportLab doesn't easily support JavaScript injection like malicious PDFs
    # This creates a clean PDF that won't trigger false positives
    c.save()
    
    print(f"Created: {filename}")
    return filename

def create_word_document(test_dir):
    """Create a Word document with normal content"""
    filename = os.path.join(test_dir, "sample_document.docx")
    
    doc = Document()
    doc.add_heading('Test Document', 0)
    
    doc.add_paragraph('This is a sample Word document created for testing the malware detection system.')
    
    doc.add_heading('Project Information', level=1)
    doc.add_paragraph('This document contains normal business content that should not trigger any malware alerts.')
    
    doc.add_paragraph('The document includes:')
    doc.add_paragraph('• Normal text content', style='List Bullet')
    doc.add_paragraph('• Standard formatting', style='List Bullet')
    doc.add_paragraph('• Common business language', style='List Bullet')
    
    doc.add_heading('Contact Information', level=1)
    doc.add_paragraph('For questions about this test, contact the development team.')
    
    doc.save(filename)
    print(f"Created: {filename}")
    return filename

def create_excel_workbook(test_dir):
    """Create an Excel workbook with formulas and data"""
    filename = os.path.join(test_dir, "sample_spreadsheet.xlsx")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Test Data"
    
    # Add headers
    ws['A1'] = 'Product'
    ws['B1'] = 'Price'
    ws['C1'] = 'Quantity'
    ws['D1'] = 'Total'
    
    # Add sample data
    products = [
        ('Widget A', 10.50, 5),
        ('Widget B', 15.75, 3),
        ('Widget C', 8.25, 7),
        ('Widget D', 12.00, 2),
        ('Widget E', 6.50, 10)
    ]
    
    for i, (product, price, qty) in enumerate(products, 2):
        ws[f'A{i}'] = product
        ws[f'B{i}'] = price
        ws[f'C{i}'] = qty
        ws[f'D{i}'] = f'=B{i}*C{i}'  # Formula to calculate total
    
    # Add sum formula
    ws['D7'] = '=SUM(D2:D6)'
    ws['A7'] = 'Grand Total'
    
    # Add another sheet with more formulas
    ws2 = wb.create_sheet("Calculations")
    ws2['A1'] = 'Advanced Formulas'
    ws2['A2'] = 'Average Price'
    ws2['B2'] = '=AVERAGE(\'Test Data\'.B2:B6)'
    ws2['A3'] = 'Max Price'
    ws2['B3'] = '=MAX(\'Test Data\'.B2:B6)'
    ws2['A4'] = 'Min Price'
    ws2['B4'] = '=MIN(\'Test Data\'.B2:B6)'
    
    wb.save(filename)
    print(f"Created: {filename}")
    return filename

def create_pdf_with_urls(test_dir):
    """Create a PDF with legitimate URLs"""
    filename = os.path.join(test_dir, "document_with_links.pdf")
    
    c = canvas.Canvas(filename, pagesize=letter)
    c.setFont("Helvetica", 12)
    
    c.drawString(100, 750, "Document with Links")
    c.drawString(100, 730, "This document contains legitimate web links:")
    c.drawString(100, 710, "• https://www.example.com")
    c.drawString(100, 690, "• https://www.google.com")
    c.drawString(100, 670, "• https://www.github.com")
    c.drawString(100, 650, "• https://www.stackoverflow.com")
    
    c.drawString(100, 600, "These are common, legitimate websites that should not")
    c.drawString(100, 580, "trigger malware alerts.")
    
    c.save()
    print(f"Created: {filename}")
    return filename

def create_empty_file(test_dir):
    """Create an empty file to test error handling"""
    filename = os.path.join(test_dir, "empty_file.pdf")
    
    with open(filename, 'wb') as f:
        pass  # Create empty file
    
    print(f"Created: {filename}")
    return filename

def create_text_file(test_dir):
    """Create a text file to test unsupported format handling"""
    filename = os.path.join(test_dir, "unsupported_format.txt")
    
    with open(filename, 'w') as f:
        f.write("This is a plain text file.\n")
        f.write("It should be rejected as an unsupported format.\n")
        f.write("The system should handle this gracefully.\n")
    
    print(f"Created: {filename}")
    return filename

def create_large_excel_file(test_dir):
    """Create a larger Excel file with more complex content"""
    filename = os.path.join(test_dir, "large_spreadsheet.xlsx")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Large Dataset"
    
    # Create headers
    headers = ['ID', 'Name', 'Category', 'Price', 'Stock', 'Revenue', 'Date']
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)
    
    # Add lots of data
    categories = ['Electronics', 'Clothing', 'Books', 'Home', 'Sports']
    for row in range(2, 1002):  # 1000 rows of data
        ws.cell(row=row, column=1, value=f'ID{row-1:04d}')
        ws.cell(row=row, column=2, value=f'Product {row-1}')
        ws.cell(row=row, column=3, value=categories[(row-2) % len(categories)])
        ws.cell(row=row, column=4, value=f'=ROUND(RAND()*100, 2)')
        ws.cell(row=row, column=5, value=f'=RANDBETWEEN(1, 100)')
        ws.cell(row=row, column=6, value=f'=D{row}*E{row}')
        ws.cell(row=row, column=7, value=f'=TODAY()-RANDBETWEEN(1, 365)')
    
    wb.save(filename)
    print(f"Created: {filename}")
    return filename

def main():
    """Create all test files"""
    print("Creating test files for Project Sentinel...")
    print("=" * 50)
    
    test_dir = create_test_directory()
    
    try:
        # Create various test files
        files_created = []
        
        files_created.append(create_simple_pdf(test_dir))
        files_created.append(create_word_document(test_dir))
        files_created.append(create_excel_workbook(test_dir))
        files_created.append(create_pdf_with_urls(test_dir))
        files_created.append(create_empty_file(test_dir))
        files_created.append(create_text_file(test_dir))
        files_created.append(create_large_excel_file(test_dir))
        
        print("=" * 50)
        print(f"Successfully created {len(files_created)} test files in '{test_dir}' directory")
        print("\nFiles created:")
        for file in files_created:
            size = os.path.getsize(file)
            print(f"  {os.path.basename(file)} ({size} bytes)")
            
    except Exception as e:
        print(f"Error creating test files: {e}")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())

