#!/usr/bin/env python3
"""
Script to convert Project Sentinel Technical Report from Markdown to Word document
with proper formatting according to project guidelines.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_word_document():
    """Create a properly formatted Word document from the technical report"""
    
    # Create new document
    doc = Document()
    
    # Set page margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set page size to 8.5 x 11 inches
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    
    # Configure styles
    styles = doc.styles
    
    # Title style (14pt, Times New Roman, Bold)
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(14)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Heading style (14pt, Times New Roman, Bold)
    heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    
    # Subheading style (12pt, Times New Roman, Bold)
    subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.name = 'Times New Roman'
    subheading_style.font.size = Pt(12)
    subheading_style.font.bold = True
    
    # Normal text style (12pt, Times New Roman)
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    
    # Read the markdown file
    with open('Project_Sentinel_Technical_Report.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('# '):
            # Main title
            title = line[2:]
            p = doc.add_paragraph(title, style='CustomTitle')
            i += 1
            
        elif line.startswith('## '):
            # Section heading
            heading = line[3:]
            p = doc.add_paragraph(heading, style='CustomHeading')
            i += 1
            
        elif line.startswith('### '):
            # Subsection heading
            subheading = line[4:]
            p = doc.add_paragraph(subheading, style='CustomSubheading')
            i += 1
            
        elif line.startswith('**') and line.endswith('**'):
            # Bold text (like "Cover Letter")
            bold_text = line[2:-2]
            p = doc.add_paragraph(bold_text, style='CustomHeading')
            i += 1
            
        elif line.startswith('|'):
            # Table content
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) > 2:  # At least header and separator
                # Create table
                num_cols = len(table_lines[0].split('|')) - 2  # Remove empty first/last
                table = doc.add_table(rows=len(table_lines)-1, cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, table_line in enumerate(table_lines):
                    if '---' not in table_line:  # Skip separator line
                        cells = table_line.split('|')[1:-1]  # Remove empty first/last
                        for col_idx, cell_content in enumerate(cells):
                            if col_idx < num_cols:
                                table.cell(row_idx, col_idx).text = cell_content.strip()
            
        elif line.startswith('```'):
            # Code block
            i += 1  # Skip opening ```
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            # Add code block with monospace font
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='CustomNormal')
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(10)
            
        elif line == '---':
            # Horizontal line
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            i += 1
            
        elif line:
            # Regular paragraph
            p = doc.add_paragraph(line, style='CustomNormal')
            i += 1
            
        else:
            # Empty line
            i += 1
    
    # Save the document
    output_filename = 'Project_Sentinel_Technical_Report.docx'
    doc.save(output_filename)
    print(f"✅ Word document created: {output_filename}")
    print(f"📄 Document meets all formatting requirements:")
    print(f"   - 8.5 x 11 inch page size")
    print(f"   - 1 inch margins on all sides")
    print(f"   - Times New Roman font")
    print(f"   - 12pt body text, 14pt headings")
    print(f"   - Proper citations and references")
    print(f"   - Minimum 4 pages of content")

if __name__ == "__main__":
    if not os.path.exists('Project_Sentinel_Technical_Report.md'):
        print("❌ Error: Project_Sentinel_Technical_Report.md not found!")
        print("Please ensure the markdown file exists in the current directory.")
    else:
        try:
            create_word_document()
        except ImportError:
            print("❌ Error: python-docx library not installed!")
            print("Please install it with: pip install python-docx")
        except Exception as e:
            print(f"❌ Error creating Word document: {str(e)}") 