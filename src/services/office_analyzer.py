import re
import os
import zipfile
from typing import Dict, List, Any
from docx import Document
from openpyxl import load_workbook
import xml.etree.ElementTree as ET

class OfficeAnalyzer:
    """
    Office document analyzer for DOCX and XLSX files
    """
    
    def __init__(self):
        # Suspicious macro patterns
        self.macro_patterns = [
            r'Auto_Open',
            r'Auto_Close',
            r'Auto_Exec',
            r'AutoExec',
            r'AutoOpen',
            r'AutoClose',
            r'Document_Open',
            r'Document_Close',
            r'Workbook_Open',
            r'Workbook_Close',
            r'Workbook_BeforeClose',
            r'Workbook_Activate',
            r'Worksheet_Activate',
            r'Worksheet_Change',
            r'Worksheet_SelectionChange'
        ]
        
        # More specific suspicious API calls and objects - avoiding Excel function names
        self.suspicious_api_patterns = [
            r'Shell\.',
            r'CreateObject\s*\(',
            r'GetObject\s*\(',
            r'WScript\.',
            r'WSCript\.',
            r'cmd\.exe',
            r'powershell\.exe',
            r'PowerShell\.exe',
            r'System\.Diagnostics\.',
            r'Process\.Start\s*\(',
            r'Application\.Run\s*\(',
            r'ExecuteExcel4Macro\s*\(',
            # More specific patterns to avoid false positives
            r'WScript\.Shell',
            r'WScript\.CreateObject',
            r'System\.IO\.File',
            r'System\.Net\.WebClient',
            r'DownloadFile\s*\(',
            r'WebClient\s*\(',
            r'HttpWebRequest',
            r'Invoke-Expression',
            r'\bIEX\b',
            r'Add-Type\s*\(',
            r'Reflection\.Assembly',
            r'FromBase64String\s*\(',
            r'Convert\.FromBase64String\s*\(',
            r'System\.Text\.Encoding',
            r'System\.Convert',
            r'RegRead\s*\(',
            r'RegWrite\s*\(',
            r'RegDelete\s*\('
        ]
        
        # Suspicious content patterns
        self.suspicious_content_patterns = [
            r'(?i)enable\s+macros?',
            r'(?i)enable\s+content',
            r'(?i)security\s+warning',
            r'(?i)protected\s+view',
            r'(?i)click\s+enable',
            r'(?i)update\s+links',
            r'(?i)external\s+data',
            r'(?i)this\s+document\s+contains\s+macros',
            r'(?i)document\s+needs\s+to\s+be\s+opened\s+in\s+compatibility\s+mode',
            r'(?i)if\s+the\s+document\s+is\s+not\s+displaying\s+correctly',
            r'(?i)please\s+enable\s+macros',
            r'(?i)for\s+correct\s+display'
        ]
        
        # More specific URL patterns - avoiding false positives
        self.url_patterns = [
            r'https?://[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}[^\s<>"\']*',  # More specific HTTP/HTTPS
            r'ftp://[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}[^\s<>"\']*',     # More specific FTP
            r'\\\\[a-zA-Z0-9.-]+\\[^\s<>"\']+',                   # UNC paths
            r'file://[a-zA-Z0-9.-]+[^\s<>"\']*',
            r'mailto:[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
            r'javascript:\s*[^\s<>"\']+',
            r'vbscript:\s*[^\s<>"\']+',
            r'data:[a-zA-Z0-9/+;=,]+',
        ]
        
        # More specific Base64 patterns - avoiding Excel formula false positives
        self.base64_patterns = [
            r'(?:^|[^a-zA-Z0-9])[A-Za-z0-9+/]{40,}={0,2}(?:[^a-zA-Z0-9]|$)',  # Longer sequences, word boundaries
            r'(?i)FromBase64String\s*\(',
            r'(?i)Convert\.FromBase64String\s*\(',
            r'(?i)System\.Convert\.FromBase64String\s*\('
        ]
        
        # Excel-specific legitimate patterns to exclude
        self.excel_legitimate_patterns = [
            r'xl/',           # Excel internal paths
            r'_rels/',        # Relationship files
            r'docProps/',     # Document properties
            r'spreadsheetml', # Excel namespace
            r'relationships', # Relationship references
            r'sharedStrings', # Shared strings
            r'styles\.xml',   # Style definitions
            r'theme\d+\.xml', # Theme files
            r'calcChain\.xml' # Calculation chain
        ]
    
    def extract_docx_content(self, file_path: str) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            print(f"DOCX text extraction error: {str(e)}")
            return ""
    
    def extract_xlsx_content(self, file_path: str) -> str:
        """Extract text content from XLSX file"""
        try:
            workbook = load_workbook(file_path, data_only=True)
            text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if cell is not None:
                            text.append(str(cell))
            return '\n'.join(text)
        except Exception as e:
            print(f"XLSX text extraction error: {str(e)}")
            return ""
    
    def is_excel_legitimate_content(self, content: str) -> bool:
        """Check if content is legitimate Excel structure"""
        for pattern in self.excel_legitimate_patterns:
            if re.search(pattern, content, re.IGNORECASE):
                return True
        return False
    
    def analyze_xml_content(self, file_path: str) -> Dict[str, Any]:
        """Analyze XML content within Office documents"""
        try:
            xml_content = ""
            xml_files = []
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                for file_name in zip_file.namelist():
                    if file_name.endswith('.xml') or file_name.endswith('.rels'):
                        try:
                            content = zip_file.read(file_name).decode('utf-8', errors='ignore')
                            xml_content += content
                            xml_files.append((file_name, content))
                        except Exception:
                            continue
            
            # Look for suspicious patterns in XML
            macro_matches = 0
            api_matches = 0
            content_matches = 0
            url_matches = 0
            base64_matches = 0
            
            found_patterns = []
            
            # Check for macro patterns
            for pattern in self.macro_patterns:
                matches = re.findall(pattern, xml_content, re.IGNORECASE)
                if matches:
                    macro_matches += len(matches)
                    found_patterns.append(f"Macro pattern: {pattern}")
            
            # Check for suspicious API patterns - but exclude legitimate Excel content
            for pattern in self.suspicious_api_patterns:
                matches = re.findall(pattern, xml_content, re.IGNORECASE)
                if matches:
                    # Filter out matches that are in legitimate Excel structures
                    valid_matches = []
                    for match in matches:
                        # Check context around the match
                        match_context = xml_content[max(0, xml_content.find(match)-100):xml_content.find(match)+100]
                        if not self.is_excel_legitimate_content(match_context):
                            valid_matches.append(match)
                    
                    if valid_matches:
                        api_matches += len(valid_matches)
                        found_patterns.append(f"Suspicious API: {pattern}")
            
            # Check for suspicious content patterns
            for pattern in self.suspicious_content_patterns:
                matches = re.findall(pattern, xml_content, re.IGNORECASE)
                if matches:
                    content_matches += len(matches)
                    found_patterns.append(f"Suspicious content: {pattern}")
            
            # Check for URLs - be more conservative
            for pattern in self.url_patterns:
                matches = re.findall(pattern, xml_content, re.IGNORECASE)
                if matches:
                    # Filter out Excel internal references
                    valid_urls = []
                    for match in matches:
                        # Skip Excel internal URLs and common false positives
                        if not any(skip in match.lower() for skip in ['schemas.openxmlformats.org', 'schemas.microsoft.com', 'w3.org', 'xml.org']):
                            valid_urls.append(match)
                    url_matches += len(valid_urls)
            
            # Check for Base64 patterns - be more conservative
            for pattern in self.base64_patterns:
                matches = re.findall(pattern, xml_content, re.IGNORECASE)
                if matches:
                    # Filter out Excel internal data
                    valid_base64 = []
                    for match in matches:
                        match_context = xml_content[max(0, xml_content.find(match)-200):xml_content.find(match)+200]
                        if not self.is_excel_legitimate_content(match_context):
                            valid_base64.append(match)
                    
                    base64_matches += len(valid_base64)
                    if len(valid_base64) > 10:  # Only flag if many suspicious base64 patterns
                        found_patterns.append(f"Base64 encoding detected")
            
            return {
                "macro_patterns": macro_matches,
                "api_patterns": api_matches,
                "content_patterns": content_matches,
                "url_patterns": url_matches,
                "base64_patterns": base64_matches,
                "found_patterns": found_patterns,
                "xml_content_length": len(xml_content)
            }
            
        except Exception as e:
            print(f"XML analysis error: {str(e)}")
            return {
                "macro_patterns": 0,
                "api_patterns": 0,
                "content_patterns": 0,
                "url_patterns": 0,
                "base64_patterns": 0,
                "found_patterns": [],
                "xml_content_length": 0
            }
    
    def count_elements(self, file_path: str) -> Dict[str, int]:
        """Count various elements in the Office document"""
        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                return {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections),
                    "total_elements": len(doc.paragraphs) + len(doc.tables)
                }
            elif file_path.endswith('.xlsx'):
                workbook = load_workbook(file_path, data_only=True)
                total_cells = 0
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value is not None:
                                total_cells += 1
                return {
                    "sheets": len(workbook.sheetnames),
                    "cells": total_cells,
                    "total_elements": total_cells
                }
            else:
                return {"total_elements": 0}
        except Exception:
            return {"total_elements": 0}
    
    def analyze_file(self, file_path: str) -> Dict[str, Any]:
        """
        Perform comprehensive Office document analysis
        Returns analysis results with threat score and confidence level
        """
        try:
            # Extract text content
            if file_path.endswith('.docx'):
                text_content = self.extract_docx_content(file_path)
            elif file_path.endswith('.xlsx'):
                text_content = self.extract_xlsx_content(file_path)
            else:
                text_content = ""
            
            text_length = len(text_content)
            
            # Analyze XML content
            xml_analysis = self.analyze_xml_content(file_path)
            
            # Count elements
            element_count = self.count_elements(file_path)
            
            # Calculate threat score (0.0 = safe, 1.0 = malicious)
            threat_score = 0.0
            reasons = []
            
            # Macro patterns (high risk)
            if xml_analysis["macro_patterns"] > 0:
                threat_score += 0.5
                reasons.append(f"Macro patterns found ({xml_analysis['macro_patterns']} instances)")
            
            # Suspicious API patterns (high risk)
            if xml_analysis["api_patterns"] > 0:
                threat_score += 0.4
                reasons.append(f"Suspicious API calls found ({xml_analysis['api_patterns']} instances)")
            
            # Suspicious content patterns (medium risk)
            if xml_analysis["content_patterns"] > 0:
                threat_score += 0.3
                reasons.append(f"Suspicious content patterns found ({xml_analysis['content_patterns']} instances)")
            
            # URL patterns (low-medium risk) - higher threshold
            if xml_analysis["url_patterns"] > 10:  # Raised threshold
                threat_score += 0.2
                reasons.append(f"Multiple external URLs found ({xml_analysis['url_patterns']} instances)")
            
            # Base64 patterns (low risk) - higher threshold
            if xml_analysis["base64_patterns"] > 20:  # Raised threshold
                threat_score += 0.1
                reasons.append(f"Excessive Base64 encoding ({xml_analysis['base64_patterns']} instances)")
            
            # Cap the threat score at 1.0
            threat_score = min(threat_score, 1.0)
            
            # Calculate confidence level (how confident we are in our analysis)
            confidence_factors = []
            base_confidence = 0.7  # Base confidence for Office documents
            
            # Increase confidence based on analysis completeness
            if xml_analysis["xml_content_length"] > 1000:
                base_confidence += 0.1
                confidence_factors.append("Complete XML analysis")
            
            if text_length > 100:
                base_confidence += 0.1
                confidence_factors.append("Text content extracted")
            
            if element_count.get("total_elements", 0) > 0:
                base_confidence += 0.1
                confidence_factors.append("Document structure analyzed")
            
            # Higher confidence if clear indicators found
            if xml_analysis["macro_patterns"] > 0 or xml_analysis["api_patterns"] > 0:
                base_confidence += 0.1
                confidence_factors.append("Strong malicious indicators")
            
            # Lower confidence for edge cases
            if xml_analysis["xml_content_length"] < 500:
                base_confidence -= 0.1
                confidence_factors.append("Limited content analysis")
            
            confidence_level = min(base_confidence, 1.0)
            
            # Generate rationale
            is_malicious = threat_score > 0.7
            if threat_score > 0:
                rationale = f"Office document analysis completed. Threat score: {threat_score:.2f}. " + "; ".join(reasons)
            else:
                rationale = "Office document analysis completed. No suspicious indicators found"
            
            # Create features dictionary
            features = {
                "suspicious_patterns": xml_analysis["macro_patterns"] + xml_analysis["api_patterns"] + xml_analysis["content_patterns"],
                "text_length": text_length,
                "urls_found": xml_analysis["url_patterns"],
                "found_patterns": xml_analysis["found_patterns"]
            }
            
            # Add document-specific features
            if file_path.endswith('.docx'):
                features["paragraph_count"] = element_count.get("paragraphs", 0)
            elif file_path.endswith('.xlsx'):
                features["cell_count"] = element_count.get("cells", 0)
                features["sheet_count"] = element_count.get("sheets", 0)
            
            return {
                "is_malicious": is_malicious,
                "threat_score": threat_score,
                "confidence_level": confidence_level,
                "confidence_factors": confidence_factors,
                "rationale": rationale,
                "features": features,
                "source": "AI Analysis"
            }
            
        except Exception as e:
            return {
                "is_malicious": False,
                "threat_score": 0.0,
                "confidence_level": 0.3,  # Low confidence due to analysis failure
                "confidence_factors": ["Analysis failed"],
                "rationale": f"Office document analysis failed: {str(e)}",
                "features": {},
                "source": "Error"
            } 