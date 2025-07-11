#!/usr/bin/env python3
"""
Advanced test file generator for Project Sentinel
Creates diverse samples to test various detection scenarios
"""

import os
import hashlib
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import struct
import random
import string

def create_large_safe_pdf():
    """Create a larger safe PDF with complex content"""
    filename = "large_safe_document.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    
    # Add multiple pages with varied content
    for page in range(5):
        c.drawString(100, 750, f"Page {page + 1} - Project Sentinel Test Document")
        c.drawString(100, 720, "This is a comprehensive test document with multiple pages.")
        c.drawString(100, 690, "It contains various text elements, formatting, and content types.")
        
        # Add some structured content
        y_pos = 650
        for i in range(10):
            c.drawString(100, y_pos, f"Line {i + 1}: Sample business content with normal text patterns.")
            y_pos -= 20
        
        # Add some numbers and data
        c.drawString(100, 400, "Financial Data:")
        c.drawString(100, 380, "Q1 Revenue: $125,000")
        c.drawString(100, 360, "Q2 Revenue: $142,500")
        c.drawString(100, 340, "Q3 Revenue: $138,750")
        
        c.showPage()
    
    c.save()
    print(f"Created: {filename}")
    return filename

def create_pdf_with_mixed_content():
    """Create PDF with some suspicious but not definitively malicious content"""
    filename = "mixed_content_document.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    
    c.drawString(100, 750, "Mixed Content Test Document")
    c.drawString(100, 720, "This document contains some patterns that might be flagged:")
    c.drawString(100, 690, "JavaScript: var x = 'normal script usage';")
    c.drawString(100, 660, "File operations: CreateFile for legitimate purposes")
    c.drawString(100, 630, "Network: http://legitimate-business-site.com")
    c.drawString(100, 600, "But overall this is a legitimate business document.")
    
    c.save()
    print(f"Created: {filename}")
    return filename

def create_heavily_suspicious_pdf():
    """Create PDF with multiple suspicious indicators"""
    filename = "heavily_suspicious_document.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    
    c.drawString(100, 750, "URGENT: Security Update Required")
    c.drawString(100, 720, "Click here to update your system immediately!")
    c.drawString(100, 690, "/JavaScript eval(unescape('%75%6E%65%73%63%61%70%65'))")
    c.drawString(100, 660, "/OpenAction /Launch cmd.exe /c powershell.exe")
    c.drawString(100, 630, "String.fromCharCode(malicious_payload_here)")
    c.drawString(100, 600, "/EmbeddedFile hidden_malware.exe")
    c.drawString(100, 570, "document.write(unescape(encoded_script))")
    c.drawString(100, 540, "/XFA forms with auto-execution")
    
    c.save()
    print(f"Created: {filename}")
    return filename

def create_complex_safe_docx():
    """Create a complex but safe Word document"""
    filename = "complex_safe_document.docx"
    doc = Document()
    
    doc.add_heading('Annual Business Report 2024', 0)
    
    # Add multiple sections
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('This comprehensive report outlines our business performance for 2024. '
                     'The document includes financial data, operational metrics, and strategic initiatives.')
    
    doc.add_heading('Financial Performance', level=1)
    doc.add_paragraph('Revenue increased by 15% year-over-year, reaching $2.5 million in total sales.')
    doc.add_paragraph('Operating expenses were well-controlled at 78% of revenue.')
    
    # Add a table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Quarter'
    hdr_cells[1].text = 'Revenue'
    hdr_cells[2].text = 'Profit'
    
    # Data rows
    data = [('Q1', '$625,000', '$125,000'),
            ('Q2', '$650,000', '$135,000'),
            ('Q3', '$675,000', '$145,000')]
    
    for i, (quarter, revenue, profit) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = quarter
        row_cells[1].text = revenue
        row_cells[2].text = profit
    
    doc.add_heading('Strategic Initiatives', level=1)
    doc.add_paragraph('Our key initiatives for the coming year include:')
    doc.add_paragraph('• Digital transformation projects')
    doc.add_paragraph('• Customer experience improvements')
    doc.add_paragraph('• Operational efficiency enhancements')
    
    doc.save(filename)
    print(f"Created: {filename}")
    return filename

def create_macro_suspicious_docx():
    """Create DOCX with macro-like suspicious content"""
    filename = "macro_suspicious_document.docx"
    doc = Document()
    
    doc.add_heading('System Configuration Document', 0)
    doc.add_paragraph('This document contains system configuration information.')
    doc.add_paragraph('Auto_Open() function detected in embedded macros')
    doc.add_paragraph('Shell.Application CreateObject("WScript.Shell") for automation')
    doc.add_paragraph('Document_Open() triggers: cmd.exe /c powershell.exe -enc')
    doc.add_paragraph('Base64 encoded payload: aGVsbG8gd29ybGQ=')
    doc.add_paragraph('Workbook_Open() macro execution enabled')
    doc.add_paragraph('CreateObject("Scripting.FileSystemObject") for file operations')
    
    doc.save(filename)
    print(f"Created: {filename}")
    return filename

def create_url_heavy_xlsx():
    """Create XLSX with many external URLs"""
    filename = "url_heavy_spreadsheet.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "External Links"
    
    ws['A1'] = "External Resources Spreadsheet"
    ws['A2'] = "This spreadsheet contains multiple external links"
    
    # Add many URLs
    urls = [
        "http://suspicious-domain.com/payload",
        "https://malware-host.net/download",
        "http://phishing-site.org/login",
        "https://fake-bank.com/secure",
        "http://trojan-download.biz/file",
        "https://ransomware-c2.net/command",
        "http://botnet-controller.org/bot",
        "https://credential-stealer.com/harvest"
    ]
    
    for i, url in enumerate(urls, 3):
        ws[f'A{i}'] = f"Link {i-2}: {url}"
    
    ws['A12'] = "Auto_Open macro reference"
    ws['A13'] = "Shell.Application automation"
    ws['A14'] = "CreateObject WScript.Shell execution"
    
    wb.save(filename)
    print(f"Created: {filename}")
    return filename

def create_formula_heavy_xlsx():
    """Create XLSX with complex formulas but safe content"""
    filename = "formula_heavy_spreadsheet.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Financial Calculations"
    
    ws['A1'] = "Financial Analysis Spreadsheet"
    ws['A2'] = "Complex calculations for business analysis"
    
    # Add headers
    ws['A4'] = "Month"
    ws['B4'] = "Revenue"
    ws['C4'] = "Expenses"
    ws['D4'] = "Profit"
    ws['E4'] = "Margin %"
    
    # Add data with formulas
    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun']
    revenues = [100000, 105000, 110000, 108000, 115000, 120000]
    expenses = [75000, 78000, 82000, 80000, 85000, 88000]
    
    for i, (month, revenue, expense) in enumerate(zip(months, revenues, expenses), 5):
        ws[f'A{i}'] = month
        ws[f'B{i}'] = revenue
        ws[f'C{i}'] = expense
        ws[f'D{i}'] = f'=B{i}-C{i}'  # Profit formula
        ws[f'E{i}'] = f'=D{i}/B{i}*100'  # Margin formula
    
    # Add summary formulas
    ws['A12'] = "Totals:"
    ws['B12'] = f'=SUM(B5:B10)'
    ws['C12'] = f'=SUM(C5:C10)'
    ws['D12'] = f'=SUM(D5:D10)'
    ws['E12'] = f'=AVERAGE(E5:E10)'
    
    wb.save(filename)
    print(f"Created: {filename}")
    return filename

def create_known_hash_files():
    """Create files with specific hashes for OSINT testing"""
    files_created = []
    
    # Create file that matches our test hash
    content = "test"  # This creates hash: 9f86d081884c7d659a2feaa0c55ad015a3bf4f1b2b0b822cd15d6c15b0f00a08
    filename = "known_malicious_hash.txt"
    with open(filename, 'w') as f:
        f.write(content)
    
    # Verify hash
    with open(filename, 'rb') as f:
        file_hash = hashlib.sha256(f.read()).hexdigest()
    
    print(f"Created: {filename} (Hash: {file_hash})")
    files_created.append(filename)
    
    # Create another test file
    content2 = "hello"  # Different hash
    filename2 = "different_hash.txt"
    with open(filename2, 'w') as f:
        f.write(content2)
    
    with open(filename2, 'rb') as f:
        file_hash2 = hashlib.sha256(f.read()).hexdigest()
    
    print(f"Created: {filename2} (Hash: {file_hash2})")
    files_created.append(filename2)
    
    return files_created

def create_edge_case_files():
    """Create edge case files for testing"""
    files_created = []
    
    # Very small PDF
    filename = "tiny_document.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    c.drawString(100, 750, "Tiny")
    c.save()
    files_created.append(filename)
    print(f"Created: {filename}")
    
    # Empty-ish DOCX
    filename = "minimal_document.docx"
    doc = Document()
    doc.add_paragraph('X')
    doc.save(filename)
    files_created.append(filename)
    print(f"Created: {filename}")
    
    # Single cell XLSX
    filename = "single_cell.xlsx"
    wb = Workbook()
    ws = wb.active
    ws['A1'] = "Single cell content"
    wb.save(filename)
    files_created.append(filename)
    print(f"Created: {filename}")
    
    return files_created

def main():
    print("Creating advanced test files for Project Sentinel...")
    
    all_files = []
    
    # Create various test files
    all_files.append(create_large_safe_pdf())
    all_files.append(create_pdf_with_mixed_content())
    all_files.append(create_heavily_suspicious_pdf())
    all_files.append(create_complex_safe_docx())
    all_files.append(create_macro_suspicious_docx())
    all_files.append(create_url_heavy_xlsx())
    all_files.append(create_formula_heavy_xlsx())
    all_files.extend(create_known_hash_files())
    all_files.extend(create_edge_case_files())
    
    print(f"\nCreated {len(all_files)} advanced test files:")
    for file in all_files:
        if os.path.exists(file):
            size = os.path.getsize(file)
            print(f"  - {file} ({size} bytes)")
    
    print("\nAdvanced test files are ready for comprehensive validation!")

if __name__ == "__main__":
    main()

