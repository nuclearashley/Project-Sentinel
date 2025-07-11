#!/usr/bin/env python3
"""
Create more realistic test samples that simulate real-world scenarios
"""

import os
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import base64

def create_phishing_pdf():
    """Create a PDF that simulates a phishing document"""
    filename = "urgent_security_update.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    
    c.drawString(100, 750, "URGENT SECURITY ALERT")
    c.drawString(100, 720, "Your system has been compromised!")
    c.drawString(100, 690, "Click the link below to secure your account immediately:")
    c.drawString(100, 660, "http://secure-bank-update.malicious-domain.com/login")
    c.drawString(100, 630, "This document contains JavaScript for automatic redirection")
    c.drawString(100, 600, "/JavaScript this.print({bUI:true,bSilent:false,bShrinkToFit:true});")
    c.drawString(100, 570, "/OpenAction << /S /JavaScript /JS (app.launchURL('http://evil.com')) >>")
    c.drawString(100, 540, "WARNING: Failure to update within 24 hours will result in account suspension")
    
    c.save()
    print(f"Created: {filename}")
    return filename

def create_invoice_scam_docx():
    """Create a DOCX that simulates an invoice scam"""
    filename = "invoice_payment_required.docx"
    doc = Document()
    
    doc.add_heading('INVOICE - PAYMENT OVERDUE', 0)
    doc.add_paragraph('Invoice #: INV-2024-7891')
    doc.add_paragraph('Amount Due: $2,847.50')
    doc.add_paragraph('Due Date: OVERDUE')
    doc.add_paragraph('')
    doc.add_paragraph('URGENT: Your payment is severely overdue. Click the link below to pay immediately:')
    doc.add_paragraph('Payment Portal: http://payment-portal.fake-company.net/pay')
    doc.add_paragraph('')
    doc.add_paragraph('This document contains macros for automatic processing:')
    doc.add_paragraph('Auto_Open() - Automatically opens payment portal')
    doc.add_paragraph('Shell.Application - Launches browser for payment')
    doc.add_paragraph('CreateObject("WScript.Shell") - System integration')
    doc.add_paragraph('Document_Open() - Triggers on document open')
    doc.add_paragraph('')
    doc.add_paragraph('Encoded payload: ' + base64.b64encode(b'malicious_script_here').decode())
    
    doc.save(filename)
    print(f"Created: {filename}")
    return filename

def create_legitimate_business_pdf():
    """Create a legitimate business document"""
    filename = "quarterly_business_report.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    
    c.drawString(100, 750, "Quarterly Business Report - Q3 2024")
    c.drawString(100, 720, "Executive Summary")
    c.drawString(100, 690, "This quarter showed strong performance across all business units.")
    c.drawString(100, 660, "Revenue increased by 12% compared to the previous quarter.")
    c.drawString(100, 630, "Key achievements include:")
    c.drawString(120, 600, "• Successful product launch in European markets")
    c.drawString(120, 580, "• Implementation of new customer service platform")
    c.drawString(120, 560, "• Expansion of development team by 25%")
    c.drawString(100, 530, "Financial Performance:")
    c.drawString(120, 500, "• Total Revenue: $4.2M")
    c.drawString(120, 480, "• Operating Expenses: $3.1M")
    c.drawString(120, 460, "• Net Profit: $1.1M")
    c.drawString(100, 430, "For detailed financial data, see attached spreadsheet.")
    c.drawString(100, 400, "Contact: finance@legitimate-company.com")
    
    c.save()
    print(f"Created: {filename}")
    return filename

def create_malware_dropper_xlsx():
    """Create XLSX that simulates a malware dropper"""
    filename = "employee_salary_data.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Salary Data"
    
    ws['A1'] = "CONFIDENTIAL - Employee Salary Information"
    ws['A2'] = "This file requires macros to be enabled for proper viewing"
    ws['A3'] = ""
    ws['A4'] = "Employee ID"
    ws['B4'] = "Name"
    ws['C4'] = "Salary"
    
    # Add some fake employee data
    employees = [
        ("EMP001", "John Smith", "$75,000"),
        ("EMP002", "Jane Doe", "$82,000"),
        ("EMP003", "Mike Johnson", "$68,000"),
    ]
    
    for i, (emp_id, name, salary) in enumerate(employees, 5):
        ws[f'A{i}'] = emp_id
        ws[f'B{i}'] = name
        ws[f'C{i}'] = salary
    
    # Add malicious content
    ws['A10'] = "Auto_Open macro enabled for data decryption"
    ws['A11'] = "Workbook_Open() - Initializes secure viewing mode"
    ws['A12'] = "Shell.Application - Required for data validation"
    ws['A13'] = "CreateObject('WScript.Shell') - Security component"
    ws['A14'] = "cmd.exe /c powershell.exe -WindowStyle Hidden -ExecutionPolicy Bypass"
    ws['A15'] = "Download URL: http://malware-host.evil.com/payload.exe"
    ws['A16'] = "Base64 payload: " + base64.b64encode(b'fake_malware_payload').decode()
    
    wb.save(filename)
    print(f"Created: {filename}")
    return filename

def create_legitimate_financial_xlsx():
    """Create a legitimate financial spreadsheet"""
    filename = "budget_analysis_2024.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Budget Analysis"
    
    ws['A1'] = "Annual Budget Analysis 2024"
    ws['A2'] = "Department Budget Allocation and Performance"
    ws['A3'] = ""
    
    # Headers
    ws['A4'] = "Department"
    ws['B4'] = "Allocated Budget"
    ws['C4'] = "Actual Spending"
    ws['D4'] = "Variance"
    ws['E4'] = "Variance %"
    
    # Data
    departments = [
        ("Marketing", 250000, 235000),
        ("Engineering", 450000, 467000),
        ("Sales", 180000, 175000),
        ("Operations", 320000, 298000),
        ("HR", 120000, 115000),
    ]
    
    for i, (dept, allocated, actual) in enumerate(departments, 5):
        ws[f'A{i}'] = dept
        ws[f'B{i}'] = allocated
        ws[f'C{i}'] = actual
        ws[f'D{i}'] = f'=C{i}-B{i}'
        ws[f'E{i}'] = f'=D{i}/B{i}*100'
    
    # Summary
    ws['A12'] = "Summary:"
    ws['B12'] = f'=SUM(B5:B9)'
    ws['C12'] = f'=SUM(C5:C9)'
    ws['D12'] = f'=SUM(D5:D9)'
    
    ws['A14'] = "Notes:"
    ws['A15'] = "• Engineering exceeded budget due to additional hiring"
    ws['A16'] = "• Marketing came in under budget with efficient campaigns"
    ws['A17'] = "• Overall budget variance is within acceptable limits"
    
    wb.save(filename)
    print(f"Created: {filename}")
    return filename

def create_resume_docx():
    """Create a legitimate resume document"""
    filename = "john_doe_resume.docx"
    doc = Document()
    
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer')
    doc.add_paragraph('Email: john.doe@email.com | Phone: (555) 123-4567')
    doc.add_paragraph('LinkedIn: linkedin.com/in/johndoe')
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('Experienced software engineer with 5+ years of experience in full-stack development. '
                     'Proficient in Python, JavaScript, and cloud technologies. Strong problem-solving skills '
                     'and experience leading development teams.')
    
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph('• Programming Languages: Python, JavaScript, Java, C++')
    doc.add_paragraph('• Frameworks: React, Django, Flask, Node.js')
    doc.add_paragraph('• Databases: PostgreSQL, MongoDB, Redis')
    doc.add_paragraph('• Cloud Platforms: AWS, Azure, Google Cloud')
    doc.add_paragraph('• Tools: Git, Docker, Kubernetes, Jenkins')
    
    doc.add_heading('Work Experience', level=1)
    doc.add_paragraph('Senior Software Engineer | TechCorp Inc. | 2021-Present')
    doc.add_paragraph('• Led development of microservices architecture serving 1M+ users')
    doc.add_paragraph('• Improved system performance by 40% through optimization')
    doc.add_paragraph('• Mentored junior developers and conducted code reviews')
    
    doc.add_paragraph('Software Engineer | StartupXYZ | 2019-2021')
    doc.add_paragraph('• Developed full-stack web applications using React and Django')
    doc.add_paragraph('• Implemented CI/CD pipelines reducing deployment time by 60%')
    doc.add_paragraph('• Collaborated with product team to define technical requirements')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    doc.add_paragraph('University of Technology | 2015-2019')
    doc.add_paragraph('GPA: 3.8/4.0')
    
    doc.save(filename)
    print(f"Created: {filename}")
    return filename

def create_contract_with_embedded_threats():
    """Create a contract document with embedded threats"""
    filename = "software_license_agreement.docx"
    doc = Document()
    
    doc.add_heading('Software License Agreement', 0)
    doc.add_paragraph('This agreement governs the use of the licensed software.')
    
    doc.add_heading('Terms and Conditions', level=1)
    doc.add_paragraph('1. The licensee agrees to the terms specified herein.')
    doc.add_paragraph('2. Software may not be redistributed without permission.')
    doc.add_paragraph('3. This agreement is effective upon installation.')
    
    # Hidden malicious content
    doc.add_paragraph('')
    doc.add_paragraph('HIDDEN SECTION - NOT VISIBLE IN NORMAL VIEW:')
    doc.add_paragraph('Auto_Open() macro will execute upon document opening')
    doc.add_paragraph('Shell.Application CreateObject("WScript.Shell")')
    doc.add_paragraph('cmd.exe /c powershell.exe -enc ' + base64.b64encode(b'malicious_powershell_command').decode())
    doc.add_paragraph('Document_Open() triggers automatic license validation')
    doc.add_paragraph('CreateObject("MSXML2.XMLHTTP") for license server communication')
    doc.add_paragraph('License server: http://license-validation.suspicious-domain.net')
    
    doc.add_heading('Contact Information', level=1)
    doc.add_paragraph('For support, contact: support@software-company.com')
    doc.add_paragraph('License validation may require internet connectivity.')
    
    doc.save(filename)
    print(f"Created: {filename}")
    return filename

def main():
    print("Creating realistic test samples for Project Sentinel...")
    
    files_created = []
    
    # Malicious/suspicious files
    files_created.append(create_phishing_pdf())
    files_created.append(create_invoice_scam_docx())
    files_created.append(create_malware_dropper_xlsx())
    files_created.append(create_contract_with_embedded_threats())
    
    # Legitimate files
    files_created.append(create_legitimate_business_pdf())
    files_created.append(create_legitimate_financial_xlsx())
    files_created.append(create_resume_docx())
    
    print(f"\nCreated {len(files_created)} realistic test files:")
    for file in files_created:
        if os.path.exists(file):
            size = os.path.getsize(file)
            print(f"  - {file} ({size:,} bytes)")
    
    print("\nRealistic test files are ready for validation!")

if __name__ == "__main__":
    main()

